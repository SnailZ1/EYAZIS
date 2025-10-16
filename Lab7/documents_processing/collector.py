from .document import Document

import os
import glob
import PyPDF2
import docx
import re

class DocumentCollector:
    """
    Класс для сбора и обработки документов из директории
    """
    
    def __init__(self, supported_extensions=None):
        if supported_extensions is None:
            self.supported_extensions = {
                '.txt': self._read_txt,
                '.pdf': self._read_pdf,
                '.docx': self._read_docx,
                '.doc': self._read_docx
            }
        else:
            self.supported_extensions = supported_extensions
        
        self.documents = []
        self.next_id = 1
    
    def _read_txt(self, file_path):
        """Чтение текстовых файлов"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='cp1251') as file:
                    return file.read()
            except:
                return ""
        except Exception as e:
            print(f"Ошибка чтения TXT файла {file_path}: {e}")
            return ""
    
    def _read_pdf(self, file_path):
        """Чтение PDF файлов"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            print(f"Ошибка чтения PDF файла {file_path}: {e}")
            return ""
    
    def _read_docx(self, file_path):
        """Чтение DOCX файлов"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            print(f"Ошибка чтения DOCX файла {file_path}: {e}")
            return ""
    
    def _get_file_title(self, file_path):
        """Извлекает заголовок из имени файла"""
        base_name = os.path.basename(file_path)
        name_without_ext = os.path.splitext(base_name)[0]
        # Заменяем подчеркивания и дефисы на пробелы
        title = re.sub(r'[_-]', ' ', name_without_ext)
        # Делаем первую букву заглавной
        return title.title()
    
    def _is_text_file(self, file_path):
        """Проверяет, является ли файл текстовым и поддерживаемым"""
        _, ext = os.path.splitext(file_path.lower())
        return ext in self.supported_extensions
    
    def collect_documents(self, directory_path, recursive=True):
        """
        Собирает все документы из указанной директории
        
        Args:
            directory_path (str): Путь к директории для сканирования
            recursive (bool): Рекурсивный поиск во вложенных папках
        
        Returns:
            list: Список объектов Document
        """
        if not os.path.exists(directory_path):
            print(f"Директория {directory_path} не существует!")
            return []
        
        print(f"Начинаем сбор документов из: {directory_path}")
        
        # Определяем метод поиска файлов
        if recursive:
            pattern = os.path.join(directory_path, "**", "*")
        else:
            pattern = os.path.join(directory_path, "*")
        
        # Собираем все файлы
        all_files = glob.glob(pattern, recursive=recursive)
        
        # Фильтруем только поддерживаемые текстовые файлы
        text_files = [f for f in all_files if os.path.isfile(f) and self._is_text_file(f)]
        
        print(f"Найдено {len(text_files)} поддерживаемых файлов")
        
        # Обрабатываем каждый файл
        for file_path in text_files:
            try:
                # Определяем расширение файла
                _, ext = os.path.splitext(file_path.lower())
                
                # Читаем содержимое файла
                content = self.supported_extensions[ext](file_path)
                
                if content and content.strip():  # Проверяем, что файл не пустой
                    # Создаем объект Document
                    title = self._get_file_title(file_path)
                    file_size = os.path.getsize(file_path)
                    
                    document = Document(
                        doc_id=self.next_id,
                        title=title,
                        content=content,
                        file_path=file_path,
                        file_type=ext.upper(),
                        file_size=file_size
                    )
                    
                    self.documents.append(document)
                    self.next_id += 1
                    
                    print(f"Обработан: {title} ({ext})")
                else:
                    print(f"Пропущен пустой файл: {file_path}")
                    
            except Exception as e:
                print(f"Ошибка обработки файла {file_path}: {e}")
        
        print(f"Сбор документов завершен. Обработано: {len(self.documents)} документов")
        return self.documents
    
    def get_documents_stats(self):
        """Возвращает статистику по собранным документам"""
        if not self.documents:
            return "Нет документов"
        
        stats = {
            'total_documents': len(self.documents),
            'total_chars': sum(len(doc.content) for doc in self.documents),
            'file_types': {},
            'avg_chars_per_doc': sum(len(doc.content) for doc in self.documents) / len(self.documents)
        }
        
        for doc in self.documents:
            stats['file_types'][doc.file_type] = stats['file_types'].get(doc.file_type, 0) + 1
        
        return stats
    
    def get_document_by_id(self, doc_id):
        """Возвращает документ по ID"""
        for doc in self.documents:
            if doc.doc_id == doc_id:
                return doc
        return None
    
    def clear_documents(self):
        """Очищает список документов"""
        self.documents = []
        self.next_id = 1